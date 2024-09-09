import sys
import docx
import re
import os
import CustomFaker
from langchain_experimental.data_anonymizer import PresidioReversibleAnonymizer
from presidio_analyzer import Pattern, PatternRecognizer
import variables

################################ CARD NUMBER ##############################
card_pattern = Pattern(
    name="card_pattern",
    regex=r"\b(?:\d[ -]*?){13,19}\b",
    score=1,
)

############################################# TIME #############################
time_pattern = Pattern(
    name="time_pattern",
    regex=r"\b(1[0-2]|0?[1-9]):[0-5][0-9]\s*(AM|PM)?\b",
    score=1,
)

######################################### Phone ###########################
phone_pattern_us = Pattern(
    name="phone_pattern_us",
    regex=r"\b(?:\+1[-.\s]?|\(?\d{3}\)?[-.\s]?)?\d{3}[-.\s]?\d{4}\b",
    score=1,
)

phone_pattern_ind = Pattern(
    name="phone_pattern_ind",
    regex=r"\b(?:\+91[-.\s]?)?[6-9]\d{9}\b",
    score=1,
)

phone_pattern_japan = Pattern(
    name="phone_pattern_japan",
    regex=r"\b(?:\+81[-.\s]?|0)?(?:70|80|90|010)[-.\s]?\d{1,4}[-.\s]?\d{4}\b",
    score=1,
)

######################################## National IDS ######################
PAN_pattern = Pattern(
    name="PAN_pattern",
    regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b",
    score=1,
)

Adhar_pattern = Pattern(
    name="Adhar_pattern",
    regex=r"\b(?:\d{4}[-\s]?){2}\d{4}\b",
    score=1,
)

########################################## Postal Code ####################################
post_pattern_us = Pattern(
    name="post_pattern_us",
    regex=r"\b\d{5}(?:-\d{4})?\b",
    score=1,
)

post_pattern_india = Pattern(
    name="post_pattern_india",
    regex=r"\b\d{6}\b",
    score=1,
)

post_pattern_japan = Pattern(
    name="post_pattern_japan",
    regex=r"\b\d{7}\b",
    score=1,
)

post_pattern_canada = Pattern(
    name="post_pattern_canada",
    regex=r"\b[A-Za-z]\d[A-Za-z] \d[A-Za-z]\d\b",
    score=1,
)

############################################# MAC Address #############################
mac_pattern = Pattern(
    name="mac_pattern",
    regex=r"\b([0-9A-Fa-f]{2}[:-]){5}[0-9A-Fa-f]{2}\b",
    score=1,
)

########## RECOGNIZERS #####

################################ CARD NUMBER ##############################
card_id_recognizer = PatternRecognizer(
    supported_entity="CARD_NUMBER",
    patterns=[card_pattern]
)

############################################# TIME #############################
time_recognizer = PatternRecognizer(
    supported_entity="TIME",
    patterns=[time_pattern]
)

######################################### Phone ###########################
phone_recognizer_us = PatternRecognizer(
    supported_entity="PHONE_NUMBER_US",
    patterns=[phone_pattern_us]
)

phone_recognizer_ind = PatternRecognizer(
    supported_entity="PHONE_NUMBER_INDIA",
    patterns=[phone_pattern_ind]
)

phone_recognizer_japan = PatternRecognizer(
    supported_entity="PHONE_NUMBER_JAPAN",
    patterns=[phone_pattern_japan]
)

######################################## National IDS ######################
pan_recognizer = PatternRecognizer(
    supported_entity="PAN_INDIA",
    patterns=[PAN_pattern]
)

adhar_recognizer = PatternRecognizer(
    supported_entity="ADHAAR_INDIA",
    patterns=[Adhar_pattern]
)

########################################## Postal Code ####################################
post_recognizer_us = PatternRecognizer(
    supported_entity="POSTAL_CODE_US",
    patterns=[post_pattern_us]
)

post_recognizer_india = PatternRecognizer(
    supported_entity="POSTAL_CODE_IND",
    patterns=[post_pattern_india]
)

post_recognizer_japan = PatternRecognizer(
    supported_entity="POSTAL_CODE_JAPAN",
    patterns=[post_pattern_japan]
)

post_recognizer_canada = PatternRecognizer(
    supported_entity="POSTAL_CODE_CANADA",
    patterns=[post_pattern_canada]
)

############################################# MAC Address #############################
mac_recognizer = PatternRecognizer(
    supported_entity="MAC_ADDRESS",
    patterns=[mac_pattern]
)


def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return ' \n '.join(full_text), doc
    except Exception as e:
        print(f"Error reading DOCX file: ")
        sys.exit(1)

def process_tag(s):
    s = s.strip('<>')
    s = re.sub(r'_\d+$', '', s)
    return s

def replace_pii(doc, old, new):
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if old in run.text:
                    # Preserve original run formatting
                    font = run.font
                    color = font.color.rgb if font.color else None
                    bold = font.bold
                    italic = font.italic
                    underline = font.underline

                    # Replace the text while preserving formatting
                    run.text = run.text.replace(old, new)

                    # Apply preserved formatting to the new text
                    if color:
                        run.font.color.rgb = color
                    run.font.bold = bold
                    run.font.italic = italic
                    run.font.underline = underline

        return doc
    
    except Exception as e:
        print(f"Error replacing PII ")
        sys.exit(1)

def modify_dict(dict):
  dict_map = {}
   
  for tag, sub_dict in dict.items():
   
    tag = process_tag(tag)
 
    if(tag == "PHONE_NUMBER_US"):
      for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_new_us_phone_number(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "PHONE_NUMBER_JAPAN"):
      for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_new_japan_phone_number(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "PHONE_NUMBER_INDIA"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_new_indian_phone_number(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "CARD_NUMBER"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_new_card_number(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "MAC_ADDRESS"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_mac(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "PAN_INDIA"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_pan_number(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "ADHAAR_INDIA"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_adhar(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "TIME"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_time(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "POSTAL_CODE_IND"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_Indian_pin(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "POSTAL_CODE_JAPAN"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_japan_pin(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "POSTAL_CODE_US"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_US_pin(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "POSTAL_CODE_CANADA"):
       for tags, old_val in sub_dict.items():
        new_val = CustomFaker.generate_random_canada_pin(old_val)
        dict_map[old_val] = new_val;
 
    if(tag == "DATE_TIME"):
       for tags, old_val in sub_dict.items():
          try:
            new_val = CustomFaker.generate_random_date(old_val)
            dict_map[old_val] = new_val
          except:
              pass
 
    if(new_val == None):
      for tags, old_val in sub_dict.items():
        new_val = old_val
        dict_map[old_val] = new_val; 
 
  return dict_map

def replace_words_in_doc(input_doc_path, output_doc_path, word_mapping):
    # Load the document
    doc = docx.Document(input_doc_path)
   
    def replace_text_in_run(run, word_mapping):
        for key, value in word_mapping.items():
            if key in run.text:
                run.text = run.text.replace(key, value)
 
    def replace_text_in_paragraph(paragraph, word_mapping):
        for run in paragraph.runs:
            replace_text_in_run(run, word_mapping)
 
    # Iterate over all paragraphs and replace words
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, word_mapping)
 
    # Iterate over all tables and replace words
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, word_mapping)
 
    # Save the modified document
    print("Now Saving... ")
    doc.save(output_doc_path)

def _divide_provided_fields(anonymizing_fields):
    default_recognizers_list = [
        card_id_recognizer,
        time_recognizer,
        phone_recognizer_us,
        phone_recognizer_ind,
        phone_recognizer_japan,
        pan_recognizer,
        post_recognizer_us,
        post_recognizer_india,
        post_recognizer_japan,
        post_recognizer_canada,
        mac_recognizer,
        adhar_recognizer
    ]

    default_anonymizer1_list = ["PERSON", "CRYPTO", "IP_ADDRESS", "EMAIL_ADDRESS", "MEDICAL_LICENSE", "US_SSN"]
    default_anonymizer2_list = ["US_PASSPORT", "DATE_TIME"]
    
    ### If no input given by user we return default lists
    if(len(anonymizing_fields) == 0):
        return default_anonymizer1_list, default_anonymizer2_list, default_recognizers_list
    
    rezognizers_list = []
    anonymizer1_list = []
    anonymizer2_list = []

    for field in anonymizing_fields:
        if(field == "CARD_NUMBER"):
            rezognizers_list.append(card_id_recognizer)
        if(field == "TIME"):
            rezognizers_list.append(time_recognizer)
        if(field == "PHONE_NUMBER_US"):
            rezognizers_list.append(phone_recognizer_us)
        if(field == "PHONE_NUMBER_INDIA"):
            rezognizers_list.append(phone_recognizer_ind)
        if(field == "PHONE_NUMBER_JAPAN"):
            rezognizers_list.append(phone_recognizer_japan)
        if(field == "POSTAL_CODE_US"):
            rezognizers_list.append(post_recognizer_us)
        if(field == "POSTAL_CODE_IND"):
            rezognizers_list.append(post_recognizer_india)
        if(field == "POSTAL_CODE_JAPAN"):
            rezognizers_list.append(post_recognizer_japan)
        if(field == "POSTAL_CODE_CANADA"):
            rezognizers_list.append(post_recognizer_canada)
        if(field == "MAC_ADDRESS"):
            rezognizers_list.append(mac_recognizer)
        if(field == "PAN_INDIA"):
            rezognizers_list.append(pan_recognizer)
        if(field == "ADHAAR_INDIA"):
            rezognizers_list.append(adhar_recognizer)
        if(field in default_anonymizer1_list):
            anonymizer1_list.append(field)
        if(field in default_anonymizer2_list):
            anonymizer2_list.append(field)
        
    return anonymizer1_list,  anonymizer2_list, rezognizers_list

# def flatten_mapping(nested_mapping):
#     flat_mapping = {}
   
#     for category, mappings in nested_mapping.items():
#         for key, value in mappings.items():
#             flat_mapping[value] = key
   
#     return flat_mapping

def get_first_name(full_name):
    # Split the full name by spaces
    parts = full_name.split()
   
    # Return the first part as the first name
    return parts[0] if parts else ''
 
#pprint.pprint(deanonymizer_mapping1)
def flatten_mapping(nested_mapping):
    flat_mapping = {}
    map = {}
 
    for category, mappings in nested_mapping.items():
        for key,value in mappings.items():
            value1 = get_first_name(value)
            if(map.get(value) is None):
              flat_mapping[value] = key
              map[value1]  = key;
            else:
              flat_mapping[value] = map[value];
           
       
    return flat_mapping

try:
    file_path = variables.input_file_path
    anonymizing_fields = variables.pii_tags
    output_file_name = variables.output_file_name
    
    # file_path = r"C:\Users\lmohan\Downloads\Vinod resume.docx"
    # anonymizing_fields = []
    # output_file_name = "check.docx"
    
    print(file_path)
    print(anonymizing_fields)
    print(output_file_name)

    text, doc = read_docx(file_path)
    print("text \n",text)
    anonymizer1_list , anonymizer2_list, recognizers_list = _divide_provided_fields(anonymizing_fields)
    print("Anon 1 list: \n", anonymizer1_list)
    print("Anon 2 list: \n", anonymizer2_list)
    # print("Anon 2 list: \n", recognizers_list)
    
except Exception as e:
    print(f"Error in Reading")
    sys.exit(51)
    
try:
    anonymizer1 = PresidioReversibleAnonymizer(analyzed_fields=anonymizer1_list)
    anonymizer2 = PresidioReversibleAnonymizer(analyzed_fields=anonymizer2_list, add_default_faker_operators=False)
    
    for rec in recognizers_list:
        anonymizer2.add_recognizer(rec)

    fake_cont1 = anonymizer1.anonymize(
        text
    ) # process fake_cont1 further to anonymize custon tags, half anonymized content
    
    fake_cont2 = anonymizer2.anonymize(
        text
    ) # we need to anonymize the text to get deanonymizer_mapping
    
    deanonymizer_mapping1 = {}
    if(len(anonymizer1_list) != 0):
        deanonymizer_mapping1 = anonymizer1.deanonymizer_mapping
        deanonymizer_mapping1 = flatten_mapping(deanonymizer_mapping1)
    
    deanonymizer_mapping2 = {}
    if(len(anonymizer2_list) != 0):
        deanonymizer_mapping2 = anonymizer2.deanonymizer_mapping
        deanonymizer_mapping2 = modify_dict(deanonymizer_mapping2) # fully anoniymized content
    
    # print("deanonymizer_mapping1: " , deanonymizer_mapping1)
    # print("deanonymizer_mapping2: " , deanonymizer_mapping2)
    
    full_deanonymizer_mapping = {**deanonymizer_mapping1, **deanonymizer_mapping2}
    print("\n full_deanonymizer_mapping: \n", full_deanonymizer_mapping)

    if os.path.exists(output_file_name):
        os.remove(output_file_name)
    
except:
    print(f"Error in final anonymizing:")
    sys.exit(52)

try:
    replace_words_in_doc(file_path, output_file_name, full_deanonymizer_mapping)

except:
    print(f"Error in replacing and saving")
    sys.exit(53)
    
# try:
#     file_name = 'check.docx'
    
#     if os.path.exists(file_name):
#         os.remove(file_name)
    
#     print("Now Saving...")
#     half_anonymized_doc.save('check.docx')
    
# except:
#     print(f"Error in Saving")
#     sys.exit(54)